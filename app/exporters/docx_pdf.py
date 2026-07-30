"""Экспорт в DOCX (python-docx) и PDF (fpdf2, шрифт с кириллицей)."""
from __future__ import annotations

from pathlib import Path

from ..config import assets_dir
from ..core.models import AppError, Transcript
from .common import group_blocks, meta_lines, ts_short


def export_docx(t: Transcript, path: str) -> None:
    try:
        import docx  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError:
        raise AppError("Не установлен python-docx. Выполните: pip install -r requirements.txt")

    doc = docx.Document()
    doc.add_heading(t.title or "Транскрипт", level=1)

    meta = doc.add_paragraph()
    run = meta.add_run("\n".join(meta_lines(t)))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for b in group_blocks(t.segments):
        p = doc.add_paragraph()
        head = p.add_run(f"[{ts_short(b.start)}] " + (f"{b.speaker}: " if b.speaker else ""))
        head.bold = True
        p.add_run(b.text)

    doc.save(path)


def export_pdf(t: Transcript, path: str) -> None:
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError:
        raise AppError("Не установлен fpdf2. Выполните: pip install -r requirements.txt")

    fonts = assets_dir() / "fonts"
    regular = fonts / "DejaVuSans.ttf"
    bold = fonts / "DejaVuSans-Bold.ttf"
    if not regular.is_file():
        raise AppError("Не найден шрифт assets/fonts/DejaVuSans.ttf для PDF.")

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.add_font("DejaVu", "", str(regular))
    pdf.add_font("DejaVu", "B", str(bold if bold.is_file() else regular))

    epw = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.set_font("DejaVu", "B", 15)
    pdf.multi_cell(epw, 8, t.title or "Транскрипт")
    pdf.ln(1)

    pdf.set_font("DejaVu", "", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(epw, 4.5, "\n".join(meta_lines(t)))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    for b in group_blocks(t.segments):
        pdf.set_font("DejaVu", "B", 10)
        head = f"[{ts_short(b.start)}]" + (f" {b.speaker}:" if b.speaker else "")
        pdf.multi_cell(epw, 5.5, head)
        pdf.set_font("DejaVu", "", 10)
        pdf.multi_cell(epw, 5.5, b.text)
        pdf.ln(2.5)

    pdf.output(path)
